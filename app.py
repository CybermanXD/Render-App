from flask import Flask, render_template, request, jsonify, send_file
import google.generativeai as genai
from google.api_core import retry
from fpdf import FPDF
from docx import Document
import re
import os

app = Flask(__name__)

# Directly store your Google API key
GOOGLE_API_KEY = 'AIzaSyAAOsCMT18wlVcgYTeRrnHPkTVLz6SrD3s'  # Replace with your actual API key
genai.configure(api_key=GOOGLE_API_KEY)

# A wrapper to handle retries
def generate_with_retry(model, prompt):
    try:
        response = model.generate_content(prompt, request_options={'retry': retry.Retry()})
        if response and response.text:
            return response.text
        else:
            raise ValueError("Empty response or text from Gemini API.")
    except Exception as e:
        print(f"Error generating content: {e}")
        return None

# Function to format content for headings and subheadings
def format_content(content):
    special_sentences = {
        "Introduction",
        "Structure",
        "Full Summary",
        "Full Analysis",
        "Literary Devices",
        "Themes and Symbols",
        "Important Questions And Answers"
    }

    # Convert text between + and + into headings with bold and larger text size
    content = re.sub(r'\+(.+?)\+', r'**\1**', content, flags=re.DOTALL)
    
    # Convert text between / and / into subheadings with bold and slightly smaller text size
    content = re.sub(r'/([^/]+?)/', r'*\1*', content, flags=re.DOTALL)

    # Bold and larger font size for special sentences
    for sentence in special_sentences:
        content = re.sub(rf'^{sentence}$', f'**{sentence}**', content, flags=re.MULTILINE)
    
    # Remove asterisks
    content = content.replace('*', '')
    
    return content

# Utility function to remove unsupported characters for PDF
def clean_text_for_pdf(text):
    return text.encode('latin-1', 'replace').decode('latin-1')

# Utility function to remove HTML tags
def remove_html_tags(text):
    clean = re.compile('<.*?>')
    return re.sub(clean, '', text)

# Generate notes by iterating over all the prompt sections
def generate_notes(novel_name, author_name):
    model = genai.GenerativeModel('gemini-1.5-flash')
    
    sections = [
        "Small Introduction of author and book ",
        "Full Structure ",
        "Detailed long Summary ",
        "Fully Detailed Long Analysis ",
        "Literary Devices Fully Explained ",
        "Themes and Symbols Fully Explained ",
        "10 very Long Important Question and Answers with numbering"
    ]
    
    custom_headings = [
        "Introduction",
        "Structure",
        "Full Summary",
        "Full Analysis",
        "Literary Devices",
        "Themes and Symbols",
        "Important Questions And Answers"
    ]
    
    full_notes = ""

    for heading, section in zip(custom_headings, sections):
        full_notes += f"**{heading}**\n\n"
        
        prompt = f"Hi Gemini, You are now an expert Professor in English Literature and expert book reviewer who can give intricate details about books, novels, short stories etc. Based on this please provide accurate and precise {section} for the {novel_name} by {author_name} without any headings. Please note: It is imperative and important to not format the text in bold. Do not give any bold text in the whole document and do not format the text in any way. Do not give any special characters. Give subheadings only inside / / without bold and give headings between + and + without bold. Do not provide any instruction from this prompt as text in the response and start the Notes directly."
        
        response_text = generate_with_retry(model, prompt)
        
        if response_text:
            formatted_text = format_content(response_text)
            formatted_text = remove_html_tags(formatted_text)
            full_notes += f"{formatted_text}\n\n"
        else:
            full_notes += f"Error generating content for section: {heading}\n\n"
    
    return full_notes

# Generate PDF with formatted content
def generate_pdf(text_file, output_file):
    pdf = FPDF()
    pdf.add_page()
    
    with open(text_file, "r", encoding="latin-1", errors="replace") as f:
        pdf.set_font("Arial", size=12)
        
        for line in f:
            cleaned_line = clean_text_for_pdf(line)
            cleaned_line = remove_html_tags(cleaned_line)
            
            if '**' in cleaned_line:
                pdf.set_font("Arial", 'B', size=16)
                cleaned_line = cleaned_line.replace('**', '')
            else:
                pdf.set_font("Arial", size=12)
            
            pdf.multi_cell(0, 10, cleaned_line)
    
    pdf.output(output_file)

# Generate DOC with formatted content
def generate_doc(text_file, output_file):
    doc = Document()
    
    with open(text_file, "r", encoding="latin-1", errors="replace") as f:
        for line in f:
            line = remove_html_tags(line)
            
            if '**' in line:
                doc.add_paragraph(line.replace('**', ''), style='Heading1')
            else:
                doc.add_paragraph(line)
    
    doc.save(output_file)

# Flask app routes
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate_notes', methods=['POST'])
def generate():
    novel_name = request.form['novelName']
    author_name = request.form['authorName']
    
    # Generate notes using multiple prompts with Gemini AI
    notes = generate_notes(novel_name, author_name)
    
    # Save the notes in a session or temporary storage for download
    with open("downloads/notes.txt", "w") as f:
        f.write(notes)
    
    return jsonify({'notes': notes})

@app.route('/download/<file_type>', methods=['GET'])
def download(file_type):
    if file_type == 'pdf':
        generate_pdf("downloads/notes.txt", "downloads/notes.pdf")
        return send_file("downloads/notes.pdf", as_attachment=True)
    elif file_type == 'doc':
        generate_doc("downloads/notes.txt", "downloads/notes.docx")
        return send_file("downloads/notes.docx", as_attachment=True)
    else:
        return "Invalid file type", 400

if __name__ == '__main__':
    if not os.path.exists('downloads'):
        os.makedirs('downloads')
    app.run(debug=True)
